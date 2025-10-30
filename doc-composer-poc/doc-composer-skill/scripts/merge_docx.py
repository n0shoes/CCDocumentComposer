#!/usr/bin/env python3
"""
Document Merger for Word Documents
Merges multiple Word documents while preserving formatting.
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

def merge_documents(master_path, document_paths, output_path):
    """
    Merge multiple Word documents into one, preserving formatting.
    
    Args:
        master_path: Path to master template document
        document_paths: List of paths to documents to merge
        output_path: Path for output document
    """
    print(f"Loading master template: {master_path}")
    
    try:
        # Open master document
        master = Document(master_path)
    except Exception as e:
        print(f"Error loading master template: {e}")
        sys.exit(1)
    
    # Track if we need a page break before first document
    # (only if master has content)
    need_break = len(master.paragraphs) > 0 or len(master.tables) > 0
    
    # Process each document
    for doc_path in document_paths:
        print(f"Merging: {doc_path}")
        
        try:
            doc = Document(doc_path)
        except Exception as e:
            print(f"Error loading document {doc_path}: {e}")
            continue
        
        # Add page break if needed (except before first doc if master is empty)
        if need_break:
            master.add_page_break()
        need_break = True  # Always add break after first document
        
        # Copy all paragraphs with their formatting
        for paragraph in doc.paragraphs:
            # Create new paragraph in master
            new_para = master.add_paragraph()
            
            # Copy paragraph properties
            new_para.alignment = paragraph.alignment
            if paragraph.style and paragraph.style.name:
                try:
                    new_para.style = paragraph.style.name
                except:
                    # Style might not exist in master, use default
                    pass
            
            # Copy paragraph format properties
            pf = paragraph.paragraph_format
            npf = new_para.paragraph_format
            
            if pf.space_before is not None:
                npf.space_before = pf.space_before
            if pf.space_after is not None:
                npf.space_after = pf.space_after
            if pf.line_spacing is not None:
                npf.line_spacing = pf.line_spacing
            if pf.left_indent is not None:
                npf.left_indent = pf.left_indent
            if pf.right_indent is not None:
                npf.right_indent = pf.right_indent
            if pf.first_line_indent is not None:
                npf.first_line_indent = pf.first_line_indent
            
            # Copy runs (text with formatting)
            for run in paragraph.runs:
                new_run = new_para.add_run(run.text)
                
                # Copy run formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.color and run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
                if run.font.highlight_color:
                    new_run.font.highlight_color = run.font.highlight_color
        
        # Copy tables
        for table in doc.tables:
            # Create new table with same dimensions
            new_table = master.add_table(
                rows=len(table.rows),
                cols=len(table.columns)
            )
            
            # Try to apply table style if it exists
            if table.style:
                try:
                    new_table.style = table.style
                except:
                    pass
            
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.rows[i].cells[j]
                    # Copy cell text (simplified - doesn't preserve all formatting)
                    new_cell.text = cell.text
                    
                    # Try to preserve cell formatting
                    if cell.paragraphs:
                        # Clear default paragraph
                        new_cell.paragraphs[0].clear()
                        
                        # Copy all paragraphs from source cell
                        for para in cell.paragraphs:
                            if new_cell.paragraphs and not new_cell.paragraphs[-1].text:
                                # Use existing empty paragraph
                                new_para = new_cell.paragraphs[-1]
                            else:
                                # Add new paragraph
                                new_para = new_cell.add_paragraph()
                            
                            # Copy paragraph text and basic formatting
                            for run in para.runs:
                                new_run = new_para.add_run(run.text)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                if run.font.size:
                                    new_run.font.size = run.font.size
        
        # Copy pictures (basic support)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                # This is simplified - full image copying requires more complex handling
                # For POC, we note that images exist but don't copy them
                print(f"  Note: Document contains images (advanced copying not implemented in POC)")
                break
    
    # Save the merged document
    print(f"Saving merged document: {output_path}")
    try:
        master.save(output_path)
        print("Merge completed successfully!")
    except Exception as e:
        print(f"Error saving document: {e}")
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(
        description="Merge Word documents while preserving formatting"
    )
    parser.add_argument(
        "--master", required=True,
        help="Path to master template document"
    )
    parser.add_argument(
        "--documents", nargs="+", required=True,
        help="Paths to documents to merge"
    )
    parser.add_argument(
        "--output", required=True,
        help="Path for output document"
    )
    
    args = parser.parse_args()
    
    # Validate inputs
    master_path = Path(args.master)
    if not master_path.exists():
        print(f"Error: Master template not found: {args.master}")
        sys.exit(1)
    
    doc_paths = []
    for doc in args.documents:
        doc_path = Path(doc)
        if not doc_path.exists():
            print(f"Warning: Document not found, skipping: {doc}")
        else:
            doc_paths.append(doc_path)
    
    if not doc_paths:
        print("Error: No valid documents to merge")
        sys.exit(1)
    
    # Perform merge
    merge_documents(str(master_path), [str(p) for p in doc_paths], args.output)

if __name__ == "__main__":
    main()
