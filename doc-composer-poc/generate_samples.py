#!/usr/bin/env python3
"""
Generate sample Word documents for the Document Composer POC
Creates both library documents and master template.
"""

import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx --break-system-packages")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def create_master_template():
    """Create the master template with corporate styling."""
    doc = Document()
    
    # Add custom styles
    styles = doc.styles
    
    # Corporate heading style
    heading_style = styles.add_style('CorpHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Add header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "ACME Corporation - Confidential"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title page content
    title = doc.add_heading('Master Document Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    para = doc.add_paragraph('This master template provides consistent formatting for all composed documents.')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    Path('master-template').mkdir(exist_ok=True)
    doc.save('master-template/corporate-template.docx')
    print("Created: master-template/corporate-template.docx")

def create_cover_page():
    """Create cover page document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Q4 2024 Quarterly Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Company name
    company = doc.add_heading('ACME Corporation', 1)
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph('December 31, 2024')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(14)
    
    # Add spacing
    for _ in range(5):
        doc.add_paragraph()
    
    # Confidentiality notice
    notice = doc.add_paragraph('CONFIDENTIAL - Internal Use Only')
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.runs[0].font.bold = True
    notice.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.save('library/cover-page.docx')
    print("Created: library/cover-page.docx")

def create_executive_summary():
    """Create executive summary document."""
    doc = Document()
    
    doc.add_heading('Executive Summary', 1)
    
    doc.add_heading('Performance Highlights', 2)
    doc.add_paragraph('• Revenue increased 15% year-over-year to $2.3B')
    doc.add_paragraph('• Operating margin improved to 18.5%')
    doc.add_paragraph('• Successful launch of three new product lines')
    doc.add_paragraph('• Expanded into two new geographic markets')
    
    doc.add_heading('Key Achievements', 2)
    doc.add_paragraph(
        'This quarter marked significant progress in our digital transformation initiative. '
        'We successfully migrated 75% of our infrastructure to the cloud, resulting in '
        '30% reduction in operational costs and improved system reliability.'
    )
    
    doc.add_heading('Strategic Focus', 2)
    doc.add_paragraph(
        'Looking ahead, we remain focused on sustainable growth through innovation '
        'and operational excellence. Our investments in R&D continue to yield positive '
        'results with a robust pipeline of products scheduled for 2025 release.'
    )
    
    doc.save('library/executive-summary.docx')
    print("Created: library/executive-summary.docx")

def create_financial_overview():
    """Create financial overview document with table."""
    doc = Document()
    
    doc.add_heading('Financial Overview', 1)
    
    doc.add_paragraph('Q4 2024 demonstrated strong financial performance across all business units.')
    
    # Add financial table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['Metric', 'Q4 2024', 'Q4 2023', 'YoY Change']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    data = [
        ['Revenue', '$2.3B', '$2.0B', '+15%'],
        ['Gross Margin', '42.5%', '40.1%', '+2.4pp'],
        ['Operating Income', '$425M', '$360M', '+18%'],
        ['Net Income', '$310M', '$265M', '+17%']
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        row = table.rows[row_idx].cells
        for col_idx, value in enumerate(row_data):
            row[col_idx].text = value
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The strong financial performance was driven by increased demand in our core markets '
        'and successful cost optimization initiatives implemented throughout the year.'
    )
    
    doc.save('library/financial-overview.docx')
    print("Created: library/financial-overview.docx")

def create_market_analysis():
    """Create market analysis document."""
    doc = Document()
    
    doc.add_heading('Market Analysis', 1)
    
    doc.add_heading('Industry Trends', 2)
    doc.add_paragraph(
        'The global market for our products continues to expand at a CAGR of 12%, '
        'driven by digital transformation initiatives and increasing automation adoption.'
    )
    
    doc.add_heading('Competitive Position', 2)
    doc.add_paragraph('• Market share increased to 23% (up from 21%)')
    doc.add_paragraph('• Ranked #2 in customer satisfaction surveys')
    doc.add_paragraph('• Leading position in emerging markets')
    
    doc.add_heading('Growth Opportunities', 2)
    doc.add_paragraph(
        'We have identified several high-growth segments that align with our core competencies. '
        'The enterprise segment shows particular promise with expected growth of 20% annually.'
    )
    
    doc.save('library/market-analysis.docx')
    print("Created: library/market-analysis.docx")

def create_risk_assessment():
    """Create risk assessment document."""
    doc = Document()
    
    doc.add_heading('Risk Assessment', 1)
    
    doc.add_heading('Key Risk Factors', 2)
    
    # Add risk table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Medium Grid 3 Accent 1'
    
    # Headers
    headers = ['Risk Category', 'Impact', 'Mitigation Strategy']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Risk data
    risks = [
        ['Supply Chain', 'Medium', 'Diversified supplier base, increased inventory'],
        ['Regulatory', 'Low', 'Compliance team, regular audits'],
        ['Market Competition', 'Medium', 'Innovation focus, customer retention programs']
    ]
    
    for row_idx, risk_data in enumerate(risks, 1):
        row = table.rows[row_idx].cells
        for col_idx, value in enumerate(risk_data):
            row[col_idx].text = value
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Our risk management framework continues to evolve with regular assessments '
        'and proactive mitigation strategies. The Risk Committee meets quarterly to '
        'review and update our risk profile.'
    )
    
    doc.save('library/risk-assessment.docx')
    print("Created: library/risk-assessment.docx")

def create_strategic_initiatives():
    """Create strategic initiatives document."""
    doc = Document()
    
    doc.add_heading('Strategic Initiatives', 1)
    
    doc.add_heading('Digital Transformation', 2)
    doc.add_paragraph(
        'Our digital transformation program is ahead of schedule with 75% of systems '
        'migrated to cloud infrastructure. This initiative has already delivered '
        '$15M in annual cost savings.'
    )
    
    doc.add_heading('Innovation Pipeline', 2)
    doc.add_paragraph('• 5 new products in final testing phase')
    doc.add_paragraph('• 12 patents filed this quarter')
    doc.add_paragraph('• R&D investment increased to 8% of revenue')
    
    doc.add_heading('Sustainability Goals', 2)
    doc.add_paragraph(
        'We remain committed to our 2030 sustainability targets. This quarter, we achieved '
        'a 10% reduction in carbon emissions and transitioned 30% of our facilities to '
        'renewable energy sources.'
    )
    
    doc.save('library/strategic-initiatives.docx')
    print("Created: library/strategic-initiatives.docx")

def create_conclusion():
    """Create conclusion document."""
    doc = Document()
    
    doc.add_heading('Conclusion', 1)
    
    doc.add_paragraph(
        'Q4 2024 represents a strong finish to an exceptional year. Our financial performance, '
        'combined with strategic progress in key initiatives, positions us well for continued '
        'growth in 2025.'
    )
    
    doc.add_heading('Looking Forward', 2)
    doc.add_paragraph(
        'As we enter 2025, we are confident in our ability to deliver value to all stakeholders. '
        'Our focus remains on sustainable growth, innovation, and operational excellence. '
        'With a strong balance sheet, talented team, and clear strategic direction, we are '
        'well-positioned to capitalize on emerging opportunities.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('Thank you for your continued support and confidence in ACME Corporation.')
    
    doc.add_paragraph()
    para = doc.add_paragraph('For questions or additional information, please contact:')
    para.runs[0].font.italic = True
    doc.add_paragraph('Investor Relations: investors@acme.com')
    doc.add_paragraph('Corporate Communications: press@acme.com')
    
    doc.save('library/conclusion.docx')
    print("Created: library/conclusion.docx")

def main():
    """Generate all sample documents."""
    print("Generating sample documents for Document Composer POC...\n")
    
    # Create directories
    Path('library').mkdir(exist_ok=True)
    Path('master-template').mkdir(exist_ok=True)
    Path('output').mkdir(exist_ok=True)
    
    # Create master template
    create_master_template()
    
    # Create library documents
    print("\nCreating library documents:")
    create_cover_page()
    create_executive_summary()
    create_financial_overview()
    create_market_analysis()
    create_risk_assessment()
    create_strategic_initiatives()
    create_conclusion()
    
    print("\n✅ All sample documents created successfully!")
    print("\nYou can now test the document composer with:")
    print("python doc-composer-skill/scripts/compose_document.py \\")
    print("    --manifest manifests/quarterly-report.md \\")
    print("    --library library/ \\")
    print("    --master master-template/corporate-template.docx \\")
    print("    --output output/Q4-2024-report.docx")

if __name__ == "__main__":
    main()
